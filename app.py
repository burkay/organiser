import base64
import os
import io
import re
import time
from datetime import datetime, timedelta
from abc import ABC, abstractmethod
import streamlit as st
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from pymongo import MongoClient
from PIL import Image
import cloudinary
import cloudinary.uploader


# ==================== DATABASE LAYER ====================

class DatabaseConnection:
    _instance = None
    _client = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    @property
    def client(self):
        if self._client is None:
            mongo_uri = st.secrets.get("MONGO_URI")
            if not mongo_uri:
                st.error("MONGO_URI secret'ı tanımlı değil.")
                st.stop()
            self._client = MongoClient(mongo_uri)
        return self._client

    @property
    def db(self):
        return self.client.get_database("organiser")


class BaseRepository(ABC):
    def __init__(self):
        self.db = DatabaseConnection().db

    @abstractmethod
    def get_collection_name(self):
        pass

    @property
    def collection(self):
        return self.db.get_collection(self.get_collection_name())


class EserlerRepository(BaseRepository):
    def get_collection_name(self):
        return "eserler"

    def insert_many(self, eserler):
        BATCH_SIZE = 5000
        for i in range(0, len(eserler), BATCH_SIZE):
            self.collection.insert_many(eserler[i:i + BATCH_SIZE])

    def search(self, query):
        # satis_fiyati UI'da gösterilmez
        return list(self.collection.find(query, {"satis_fiyati": 0}))

    def get_distinct_sanatcilar(self):
        return sorted(self.collection.distinct("sanatci", {"sanatci": {"$ne": ""}}))


class AyarlarRepository(BaseRepository):
    def get_collection_name(self):
        return "ayarlar"

    def get_access_code(self):
        ayar = self.collection.find_one({"tip": "giris_kontrol"})
        if not ayar or "sifre" not in ayar:
            raise ValueError("Giriş şifresi MongoDB'de tanımlı değil.")
        return ayar["sifre"]


class LogRepository(BaseRepository):
    def get_collection_name(self):
        return "ziyaretci_loglari"

    def log_login_attempt(self, entered_code, success):
        try:
            self.collection.insert_one({
                "ip_adresi":     self._get_ip_address(),
                "girilen_sifre": entered_code,
                "basarili":      success,
                "tarih_saat":    datetime.now(),
                "session_id":    self._get_session_id(),
            })
        except Exception:
            pass

    @staticmethod
    def _get_ip_address():
        try:
            if hasattr(st, 'context') and hasattr(st.context, 'headers'):
                ip = st.context.headers.get("X-Forwarded-For", "unknown")
                if ip == "unknown":
                    ip = st.context.headers.get("X-Real-IP", "unknown")
                return ip
        except Exception:
            pass
        return "unknown"

    @staticmethod
    def _get_session_id():
        try:
            from streamlit.runtime.scriptrunner import get_script_run_ctx
            ctx = get_script_run_ctx()
            return ctx.session_id if ctx else "unknown"
        except Exception:
            return "unknown"


# ==================== BUSINESS LOGIC LAYER ====================

class AuthenticationService:
    def __init__(self):
        self.ayarlar_repo = AyarlarRepository()
        self.log_repo = LogRepository()

    def verify_code(self, code):
        if not code:
            return False
        try:
            correct_code = self.ayarlar_repo.get_access_code()
            is_correct = code == correct_code
            self.log_repo.log_login_attempt(code, is_correct)
            return is_correct
        except Exception as e:
            st.error(f"Doğrulama hatası: {e}")
            return False


class SessionManager:
    TIMEOUT_HOURS = 1

    @staticmethod
    def initialize():
        if "authenticated" not in st.session_state:
            st.session_state.authenticated = False
        if "login_time" not in st.session_state:
            st.session_state.login_time = None

    @staticmethod
    def is_authenticated():
        return st.session_state.get("authenticated", False)

    @classmethod
    def check_timeout(cls):
        if cls.is_authenticated() and st.session_state.login_time:
            if datetime.now() - st.session_state.login_time > timedelta(hours=cls.TIMEOUT_HOURS):
                cls.logout()
                st.warning("Oturum süresi doldu. Lütfen tekrar giriş yapın.")
                st.rerun()

    @staticmethod
    def login():
        st.session_state.authenticated = True
        st.session_state.login_time = datetime.now()

    @staticmethod
    def logout():
        st.session_state.authenticated = False
        st.session_state.login_time = None

    @classmethod
    def get_remaining_time(cls):
        if not cls.is_authenticated() or not st.session_state.login_time:
            return timedelta(0)
        return timedelta(hours=cls.TIMEOUT_HOURS) - (datetime.now() - st.session_state.login_time)


class CloudinaryService:
    """
    Cloudinary görsel yükleme servisi.
    Secrets: CLOUDINARY_CLOUD_NAME, CLOUDINARY_API_KEY, CLOUDINARY_API_SECRET
    """
    _configured = False

    @classmethod
    def _configure(cls):
        if cls._configured:
            return
        cloudinary.config(
            cloud_name = st.secrets["CLOUDINARY_CLOUD_NAME"],
            api_key    = st.secrets["CLOUDINARY_API_KEY"],
            api_secret = st.secrets["CLOUDINARY_API_SECRET"],
            secure     = True,
        )
        cls._configured = True

    @classmethod
    def upload(cls, image_bytes: bytes, public_id: str) -> str:
        """
        Görsel byte'larını Cloudinary'e yükle, URL döndür.
        public_id → tekrar yüklenirse üzerine yazar (idempotent).
        """
        cls._configure()
        result = cloudinary.uploader.upload(
            image_bytes,
            public_id      = public_id,
            overwrite      = True,
            resource_type  = "image",
            folder         = "muzayede",
            transformation = [{"width": 800, "crop": "limit", "quality": "auto"}],
        )
        return result["secure_url"]


class MuzayedeParser:
    """
    Müzayede kataloğu .docx parse işlemleri.

    Her eser bloğunun yapısı (XML sırası):
      <p>  →  Görsel (w:drawing içeren paragraf)
      <p>  →  Galeri / Sahip  (görsel hemen arkasındaki ilk dolu satır)
      <p>  →  Sanatçı + yıl
      <p>  →  Eser adı
      <p>  →  Teknik detaylar
      <p>  →  (opsiyonel) Satış fiyatı  →  sadece DB'ye, UI'da gizli

    Sahip satırı ayrıca _is_sahip testi yapılmaz;
    görsel paragrafının hemen arkasından gelen ilk dolu satır sahip olarak alınır.
    Bu yaklaşım tüm format varyasyonlarını kapsar.
    """

    _BLIP_ATTR = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'

    @staticmethod
    def _is_fiyat(text: str) -> bool:
        return bool(re.search(r'\d[\d\.,]+\s*(TL|₺)', text, re.IGNORECASE))

    @classmethod
    def _extract_image_bytes(cls, para_elem, doc_part):
        """Paragraf elementinden embed görsel byte'larını çıkar."""
        blips = para_elem.findall('.//' + qn('a:blip'))
        if not blips:
            return None
        rId = blips[0].get(cls._BLIP_ATTR)
        if not rId or rId not in doc_part.rels:
            return None
        try:
            return doc_part.rels[rId].target_part.blob
        except Exception:
            return None

    @classmethod
    def parse(cls, doc: Document, upload_images: bool = False,
              progress_callback=None) -> list:
        """
        Document nesnesini parse et; eser listesi döndür.

        Görsel paragrafı → hemen arkasındaki ilk dolu satır sahip olarak alınır.
        Bu yaklaşım sahip satırındaki tüm format farklılıklarını kapsar.

        upload_images=True    → görselleri Cloudinary'e yükler ve gorsel_url ekler.
        upload_images=False   → gorsel_url boş kalır (hızlı önizleme).
        progress_callback     → her eser işlenince callback(done, total) çağrılır.
                                 İlk geçişte total bilinmediği için None geçilebilir.
        """
        body_children = list(doc.element.body)
        doc_part      = doc.part

        nodes = []
        for child in body_children:
            texts  = child.findall('.//' + qn('w:t'))
            text   = ''.join(t.text or '' for t in texts).strip()
            is_img = bool(child.findall('.//' + qn('w:drawing')))
            nodes.append({"elem": child, "text": text, "is_img": is_img})

        # Toplam eser sayısını önceden hesapla (progress için)
        toplam_eser = sum(
            1 for idx, nd in enumerate(nodes)
            if nd["is_img"] and idx + 1 < len(nodes)
        )

        artworks    = []
        lot_counter = 0
        i           = 0

        while i < len(nodes):
            node = nodes[i]

            if not node["is_img"]:
                i += 1
                continue

            img_elem = node["elem"]

            # Görsel sonrası ilk anlamlı node'u bul
            j = i + 1
            while j < len(nodes) and not nodes[j]["text"] and not nodes[j]["is_img"]:
                j += 1

            # Sonraki anlamlı node başka bir görsel ise bu görsel başlıksız, atla
            if j >= len(nodes) or nodes[j]["is_img"]:
                i += 1
                continue

            lot_counter += 1
            sahip = nodes[j]["text"]

            # Devamındaki satırları topla (bir sonraki görsele kadar)
            lines = []
            k = j + 1
            while k < len(nodes) and len(lines) < 6:
                if nodes[k]["is_img"]:
                    break
                t = nodes[k]["text"]
                if t:
                    lines.append(t)
                else:
                    if k + 1 < len(nodes) and not nodes[k + 1]["text"]:
                        break
                k += 1

            sanatci  = lines[0] if len(lines) > 0 else ""
            eser_adi = lines[1] if len(lines) > 1 else ""
            detay    = lines[2] if len(lines) > 2 else ""

            satis_fiyati = ""
            for ln in reversed(lines[2:]):
                if cls._is_fiyat(ln):
                    satis_fiyati = ln
                    break

            gorsel_url = ""
            if upload_images:
                img_bytes = cls._extract_image_bytes(img_elem, doc_part)
                if img_bytes:
                    public_id = f"lot_{lot_counter}"
                    try:
                        gorsel_url = CloudinaryService.upload(img_bytes, public_id)
                    except Exception as e:
                        st.warning(f"Lot {lot_counter} görseli yüklenemedi: {e}")

            artworks.append({
                "lot_no":       lot_counter,
                "sahip":        sahip,
                "sanatci":      sanatci,
                "eser_adi":     eser_adi,
                "detay":        detay,
                "gorsel_url":   gorsel_url,
                "satis_fiyati": satis_fiyati,
            })

            if progress_callback:
                progress_callback(lot_counter, toplam_eser)

            i = k

        return artworks

# ==================== PRESENTATION LAYER ====================

class LoginView:
    def __init__(self, auth_service):
        self.auth_service = auth_service

    def render(self):
        # Logo varsa st.image ile göster (unsafe_allow_html gerektirmez)
        logo_path = "logo.png"
        if os.path.exists(logo_path):
            try:
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    st.image(logo_path, width='stretch')
            except Exception:
                pass

        st.title("🔐 Giriş")
        st.markdown("---")

        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.subheader("Lütfen erişim kodunu girin")
            with st.form(key="login_form", clear_on_submit=False):
                access_code = st.text_input(
                    "Erişim Kodu",
                    type="password",
                    placeholder="Erişim kodunu girin ve Enter'a basın...",
                    key="access_code_input"
                )
                submitted = st.form_submit_button("Giriş Yap", width='stretch')
                if submitted and access_code:
                    if self.auth_service.verify_code(access_code):
                        SessionManager.login()
                        st.success("✅ Giriş başarılı!")
                        time.sleep(0.5)
                        st.rerun()
                    else:
                        st.error("❌ Hatalı erişim kodu!")
            st.caption("Erişim kodunu bilmiyorsanız, lütfen yönetici ile iletişime geçin.")


class MainView:
    GOSTERIM_LIMITI = 2000
    KART_KOLONLARI  = 4   # Arama sonuçlarında yan yana kaç kart

    def __init__(self, eserler_repo):
        self.eserler_repo = eserler_repo

    def render(self):
        self._render_header()
        # Yükleme aktifse önce işlemi tamamla, sonra normal UI'ı göster
        if st.session_state.get("yukleniyor", False):
            self._render_sidebar()
            self._do_upload()
        else:
            self._render_sidebar()
        self._render_search()

    def _render_header(self):
        col1, col2, col3 = st.columns([6, 1, 1])
        with col2:
            remaining = SessionManager.get_remaining_time()
            st.caption(f"⏱️ Kalan: {int(remaining.total_seconds() // 60)} dk")
        with col3:
            if st.button("🚪 Çıkış"):
                SessionManager.logout()
                st.rerun()
        st.title("🏛️")

    def _render_sidebar(self):
        st.sidebar.header("📤 Eser Dosyası Yükleme")
        st.sidebar.caption(
            "Müzayede kataloğu .docx dosyası yükleyin. "
            "Her sayfa bir eser: Sahip · Sanatçı · Eser Adı · Detay · (Fiyat)"
        )

        # Yükleme devam ediyorsa dosya seçimi ve checkbox deaktif
        yukleniyor = st.session_state.get("yukleniyor", False)

        uploaded_file = st.sidebar.file_uploader(
            "Word dosyası seçin (.docx)",
            type=["docx"],
            help="Sadece .docx formatı kabul edilir.",
            disabled=yukleniyor,
        )

        gorsel_yukle = st.sidebar.checkbox(
            "Görselleri yükle", value=True, disabled=yukleniyor
        )

        if uploaded_file and not yukleniyor:
            self._handle_file_upload(uploaded_file, gorsel_yukle)

        if yukleniyor:
            st.sidebar.info("⏳ Yükleme devam ediyor, lütfen bekleyin...")

    def _handle_file_upload(self, uploaded_file, gorsel_yukle: bool):
        try:
            file_key = f"docx_bytes_{uploaded_file.name}"
            if file_key not in st.session_state:
                st.session_state[file_key] = uploaded_file.read()

            file_bytes = io.BytesIO(st.session_state[file_key])
            doc = Document(file_bytes)
            kayitlar = MuzayedeParser.parse(doc, upload_images=False)

            if not kayitlar:
                st.sidebar.warning(
                    "Bu dosyada geçerli eser bloğu bulunamadı. "
                    "Her eser bir Sahip/Galeri satırıyla başlamalı."
                )
                return

            st.sidebar.success(
                f"Toplam **{len(kayitlar)}** eser bulundu. "
                f"Eklemek için butona tıklayın."
            )

            if st.sidebar.button("Eserleri Veritabanına Ekle", disabled=False):
                st.session_state["yukleniyor"] = True
                st.rerun()

        except Exception as e:
            st.sidebar.error(f"Dosya okuma hatası: {e}")

    def _do_upload(self):
        """Yükleme işlemini gerçekleştir — yukleniyor=True olduğunda çağrılır."""
        # Hangi dosya key'i var?
        file_key = next(
            (k for k in st.session_state if k.startswith("docx_bytes_")), None
        )
        if not file_key:
            st.session_state["yukleniyor"] = False
            return

        dosya_adi = file_key.replace("docx_bytes_", "")
        gorsel_yukle = st.session_state.get("gorsel_yukle_tercih", True)

        try:
            t_baslangic = time.perf_counter()
            with st.sidebar:
                fresh_bytes = io.BytesIO(st.session_state[file_key])
                doc = Document(fresh_bytes)

                kayitlar_on = MuzayedeParser.parse(doc, upload_images=False)
                toplam = len(kayitlar_on)

                st.markdown("**Eserler yükleniyor...**")
                progress_bar = st.progress(0)
                durum_yazisi = st.empty()

                fresh_bytes2 = io.BytesIO(st.session_state[file_key])
                doc2 = Document(fresh_bytes2)
                kayitlar = MuzayedeParser.parse(
                    doc2,
                    upload_images=gorsel_yukle,
                    progress_callback=lambda done, total: (
                        progress_bar.progress(done / total),
                        durum_yazisi.caption(f"{done} / {total} eser işlendi")
                    )
                )
                progress_bar.progress(1.0)
                durum_yazisi.caption(f"{toplam} / {toplam} eser işlendi")

            for k in kayitlar:
                k["dosya_adi"] = dosya_adi

            self.eserler_repo.insert_many(kayitlar)
            sure_toplam = time.perf_counter() - t_baslangic

            st.sidebar.success(f"✅ {len(kayitlar)} eser {sure_toplam:.2f} sn'de eklendi.")
            del st.session_state[file_key]

        except Exception as e:
            st.sidebar.error(f"Hata: {e}")
        finally:
            st.session_state["yukleniyor"] = False
            st.rerun()

    def _render_search(self):
        col1, col2, col3 = st.columns([3, 1, 1])
        with col1:
            search_query = st.text_input(
                "Anahtar kelime (eser adı, sanatçı, sahip, detay)",
                placeholder="Örn. Ertuğrul Ateş, yağlıboya, Levent Gürel..."
            )
        with col2:
            lot_no_query = st.text_input("Lot No", placeholder="Örn. 37")
        with col3:
            sanatci_liste = [""] + self.eserler_repo.get_distinct_sanatcilar()
            sanatci_filtre = st.selectbox("Sanatçıya göre filtrele", sanatci_liste)

        sorgu = self._build_query(search_query, lot_no_query, sanatci_filtre)
        self._show_results(sorgu)

    def _build_query(self, search_query, lot_no_query, sanatci_filtre):
        sorgu = {}

        if lot_no_query.strip():
            try:
                sorgu["lot_no"] = int(lot_no_query.strip())
            except ValueError:
                pass

        if search_query:
            sorgu["$or"] = [
                {"eser_adi": {"$regex": search_query, "$options": "i"}},
                {"sanatci":  {"$regex": search_query, "$options": "i"}},
                {"sahip":    {"$regex": search_query, "$options": "i"}},
                {"detay":    {"$regex": search_query, "$options": "i"}},
            ]

        if sanatci_filtre:
            sorgu["sanatci"] = sanatci_filtre

        return sorgu

    def _show_results(self, sorgu):
        try:
            items = self.eserler_repo.search(sorgu)
        except Exception as e:
            st.error(f"Veritabanı hatası: {e}")
            items = []

        toplam = len(items)
        st.subheader(f"🔍 {toplam:,} Eserde Ara" if toplam else "🔍 Eserlerde Ara")

        if items:
            self._display_results(items)
        else:
            st.info(
                "Sonuç bulunamadı. Sol taraftan .docx dosyası yükleyip "
                "'Eserleri Veritabanına Ekle' ile havuzu doldurun."
            )

    def _display_results(self, items):
        toplam = len(items)
        if toplam > self.GOSTERIM_LIMITI:
            st.info(f"İlk **{self.GOSTERIM_LIMITI}** kayıt gösteriliyor (toplam {toplam}).")
            items = items[:self.GOSTERIM_LIMITI]

        # Seçili eser varsa dialog aç
        if st.session_state.get("secili_eser") is not None:
            self._render_dialog(st.session_state["secili_eser"])

        self._render_list(items)

    @st.dialog("Eser Detayı", width="large")
    def _render_dialog(self, item):
        """st.dialog ile native modal — kapat butonu otomatik gelir."""
        gorsel_url = item.get("gorsel_url", "")
        lot  = item.get("lot_no", "")
        ad   = item.get("eser_adi") or "—"
        san  = item.get("sanatci") or "—"
        sah  = item.get("sahip") or "—"
        det  = item.get("detay") or "—"
        dosya = item.get("dosya_adi") or ""

        if gorsel_url:
            st.markdown(
                f"<img src='{gorsel_url}' style='max-width:100%;border-radius:8px;"
                f"display:block;margin:0 auto 1rem;'/>",
                unsafe_allow_html=True
            )

        st.markdown(f"### Lot {lot} · {ad}")
        st.markdown(f"*{san}*")
        st.divider()
        col1, col2 = st.columns([1, 3])
        col1.markdown("**Sahip**")
        col2.markdown(sah)
        col1.markdown("**Detay**")
        col2.markdown(det)
        col1.markdown("**Dosya**")
        col2.markdown(f"<small style='color:#888'>{dosya}</small>", unsafe_allow_html=True)

        if st.button("Kapat", use_container_width=False):
            st.session_state["secili_eser"] = None
            st.rerun()

    def _render_list(self, items):
        """Satır satır liste görünümü."""
        h1, h2, h3, h4, h5 = st.columns([1, 4, 3, 3, 2])
        h1.markdown("**Lot**")
        h2.markdown("**Eser Adı**")
        h3.markdown("**Sanatçı**")
        h4.markdown("**Sahip**")
        h5.markdown("")
        st.divider()

        for idx, item in enumerate(items):
            c1, c2, c3, c4, c5 = st.columns([1, 4, 3, 3, 2])
            c1.write(item.get("lot_no", ""))
            c2.write(item.get("eser_adi") or "—")
            c3.write(item.get("sanatci") or "—")
            c4.write(item.get("sahip") or "—")
            if c5.button("Detay", key=f"detay_{idx}"):
                st.session_state["secili_eser"] = item
                st.rerun()


# ==================== APPLICATION ====================

class Application:
    def __init__(self):
        self._setup_page()
        SessionManager.initialize()
        self.auth_service = AuthenticationService()
        self.eserler_repo = EserlerRepository()
        self.login_view = LoginView(self.auth_service)
        self.main_view = MainView(self.eserler_repo)

    @staticmethod
    def _setup_page():
        st.set_page_config(
            page_title="SeleSys",
            layout="wide",
            page_icon="favicon.png" if os.path.exists("favicon.png") else "logo.png",
        )

    def run(self):
        SessionManager.check_timeout()
        if SessionManager.is_authenticated():
            self.main_view.render()
        else:
            self.login_view.render()


# ==================== ENTRY POINT ====================

if __name__ == "__main__":
    app = Application()
    app.run()
